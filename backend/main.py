# backend/main.py
import os
import io
import json
import hashlib
from typing import Optional, List, Dict, Any
from datetime import datetime, timedelta
import re

from fastapi.responses import StreamingResponse
from docx import Document
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from jose import jwt, JWTError

import fitz  # PyMuPDF
import docx
from Crypto.Cipher import AES
from Crypto.Random import get_random_bytes
from base64 import b64encode, b64decode
from dotenv import load_dotenv

# AI models
from ai_models.indian_legal_bert import IndianLegalBERT
from ai_models.risk_engine import AdvancedRiskEngine
from utils.file_processor import FileProcessor

def build_corrected_document(clauses, analysis_results, threshold=0.20):
    """
    Reconstruct the full document by replacing risky clauses
    with their improved versions while preserving formatting.
    """

    corrected_parts = []

    # Create a mapping from clause_text → improved_clause
    improved_map = {}
    for r in analysis_results:
        score = float(r.get("risk_score", 0.0))
        if score >= threshold and r.get("improved_clause"):
            improved_map[r["clause_text"].strip()] = r["improved_clause"]

    for original in clauses:
        stripped = original.strip()

        if stripped in improved_map:
            # Insert corrected version (same formatting as original)
            corrected_parts.append(improved_map[stripped])
        else:
            # Keep original clause untouched
            corrected_parts.append(original)

    # Join using two newlines (paragraph-level formatting preserved)
    final_doc = "\n\n".join(corrected_parts)
    return final_doc


def build_corrected_docx(corrected_text: str) -> bytes:
    """
    Creates a clean DOCX containing ONLY the corrected (AI-improved) document.
    """
    doc = Document()

    # Title
    title = doc.add_heading("AI-Corrected Legal Document", level=1)
    title.alignment = 1  # center

    # Corrected document section
    doc.add_heading("Corrected Document", level=2)

    # Add corrected text paragraph-by-paragraph
    for para in corrected_text.split("\n"):
        doc.add_paragraph(para)

    # Save to memory
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==============================
# Load env
# ==============================
load_dotenv()

SECRET_KEY = os.getenv("SECRET_KEY", "dev-secret")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = int(os.getenv("ACCESS_TOKEN_EXPIRE_MINUTES", "60"))
FRONTEND_URL = os.getenv("FRONTEND_URL", "http://localhost:3000")
UPLOAD_DIR = "./uploads"
AES_KEY_B64 = os.getenv("AES_KEY_B64")
SINGLE_USER_EMAIL = os.getenv("SINGLE_USER_EMAIL", "admin@example.com")
SINGLE_USER_PASSWORD = os.getenv("SINGLE_USER_PASSWORD", "changeme")

os.makedirs(UPLOAD_DIR, exist_ok=True)

# ==============================
# Model Initialization
# ==============================
try:
    legal_bert = IndianLegalBERT()
    risk_engine = AdvancedRiskEngine(db_session=None)
    print("[DEBUG] AI Models Loaded")
except Exception as e:
    print("[ERROR] Model failed:", e)

    class _BrokenModel:
        def __getattr__(self, name):
            def _err(*args, **kwargs):
                raise RuntimeError(f"Model {name} unavailable: {e}")
            return _err

    legal_bert = _BrokenModel()
    risk_engine = _BrokenModel()

# ==============================
# FastAPI App
# ==============================
app = FastAPI(title="Legal AI Analyzer")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   # Allow all for debugging — change later
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==============================
# AES Helpers
# ==============================
def _ensure_aes_key():
    if AES_KEY_B64:
        key = b64decode(AES_KEY_B64)
        if len(key) in (16, 24, 32):
            return key
    return hashlib.sha256(SECRET_KEY.encode()).digest()[:32]

AES_KEY = _ensure_aes_key()

def _pad(b):
    pad_len = AES.block_size - (len(b) % AES.block_size)
    return b + bytes([pad_len]) * pad_len

def _unpad(b):
    return b[:-b[-1]]

def encrypt_bytes(data):
    iv = get_random_bytes(16)
    cipher = AES.new(AES_KEY, AES.MODE_CBC, iv)
    return b64encode(iv + cipher.encrypt(_pad(data))).decode()

def decrypt_bytes(b64str):
    raw = b64decode(b64str)
    iv, ct = raw[:16], raw[16:]
    cipher = AES.new(AES_KEY, AES.MODE_CBC, iv)
    return _unpad(cipher.decrypt(ct))

# ==============================
# Auth Models
# ==============================
class Token(BaseModel):
    access_token: str
    token_type: str

class LoginRequest(BaseModel):
    email: str
    password: str

def create_access_token(data):
    expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    return jwt.encode({**data, "exp": expire}, SECRET_KEY, ALGORITHM)

async def get_current_user(request: Request):
    auth = request.headers.get("Authorization", "")
    if not auth.startswith("Bearer "):
        raise HTTPException(401, "Missing token")
    token = auth.split()[1]
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        if payload["sub"] != SINGLE_USER_EMAIL:
            raise Exception()
        return {"email": payload["sub"]}
    except:
        raise HTTPException(401, "Invalid token")

# ==============================
# File extraction helpers
# ==============================
def extract_text_from_pdf_bytes(b):
    try:
        text = []
        with fitz.open(stream=b, filetype="pdf") as doc:
            for page in doc:
                text.append(page.get_text("text"))
        return "\n".join(text)
    except:
        return b.decode(errors="ignore")

def extract_text_from_docx_bytes(b):
    try:
        d = docx.Document(io.BytesIO(b))
        return "\n".join(p.text for p in d.paragraphs)
    except:
        return b.decode(errors="ignore")

# ==============================
# Improved Clause Generator
# ==============================
def generate_improved_clause(clause_text: str, clause_type: str):
    """
    Heuristic + model-based rewrite of risky clauses.
    Ensures ALWAYS returns a safer rewritten clause.
    """
    t = clause_text.strip()

    # --- heuristic fixes ---
    safe = re.sub(r'unlimited liability', "liability capped to reasonable amount", t, flags=re.I)
    safe = re.sub(r'50% per month', "2% per month (reasonable statutory limit)", safe, flags=re.I)
    safe = re.sub(r'non[- ]?compete.*?(\d+)\s*years?',
                  "non-compete capped at 2 years with clearly defined scope",
                  safe, flags=re.I)

    if safe != t:
        return safe.rstrip(".") + "."

    # --- model rewrite fallback ---
    try:
        prompt = (
            "Rewrite this legal clause into a safer, compliant lower-risk version "
            "following Indian Contract Act principles. Keep meaning, remove harsh penalties, "
            "replace vagueness with concrete limits:\n\n"
            f"{t}"
        )
        out = legal_bert.summarizer(prompt, max_length=160, truncation=True)
        if isinstance(out, list) and out[0].get("generated_text"):
            rewritten = out[0]["generated_text"].strip()
            return rewritten.rstrip(".") + "."
    except:
        pass

    # --- ultimate fallback ---
    return (
        "Revised safer clause: Obligations must comply with applicable Indian laws. "
        "Liability is limited to direct damages only, capped reasonably, and timelines "
        "should be explicitly defined (e.g., 30 days)."
    )

# ==============================
# /login
# ==============================
@app.post("/login", response_model=Token)
def login(body: LoginRequest):
    if body.email != SINGLE_USER_EMAIL or body.password != SINGLE_USER_PASSWORD:
        raise HTTPException(401, "Invalid login")
    return {"access_token": create_access_token({"sub": body.email}), "token_type": "bearer"}

# ==============================
# /upload
# ==============================
@app.post("/upload")
async def upload_file(file: UploadFile = File(...), user=Depends(get_current_user)):
    content = await file.read()
    encrypted = encrypt_bytes(content)

    stamp = datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
    digest = hashlib.sha256(content).hexdigest()[:12]
    fname = f"{stamp}_{digest}_{file.filename}.enc"

    with open(os.path.join(UPLOAD_DIR, fname), "w") as f:
        f.write(encrypted)

    return {"stored_filename": fname}

# ==============================
# /analyze (Part 1 ends here)
# ==============================

# ==============================
# /analyze endpoint (continued)
# ==============================

@app.post("/analyze")
async def analyze_document(payload: dict, user=Depends(get_current_user)):
    """
    Analyze uploaded text or stored encrypted file and return structured results.
    Payload: { "text": "...", "stored_filename": "..." }
    """
    text = payload.get("text")
    stored_filename = payload.get("stored_filename")

    # --- If stored file provided, decrypt & extract text ---
    if not text and stored_filename:
        path = os.path.join(UPLOAD_DIR, stored_filename)
        if not os.path.exists(path):
            raise HTTPException(status_code=404, detail="stored file not found")
        enc = open(path, "r", encoding="utf-8").read()
        try:
            raw = decrypt_bytes(enc)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to decrypt stored file: {e}")

        lower = path.lower()
        if lower.endswith(".docx.enc"):
            text = extract_text_from_docx_bytes(raw)
        else:
            # try PDF first; if empty, fallback to docx / raw decode
            text = extract_text_from_pdf_bytes(raw) or extract_text_from_docx_bytes(raw) or (raw.decode("utf-8", errors="replace") if isinstance(raw, (bytes, bytearray)) else str(raw))

    if not text or not text.strip():
        raise HTTPException(status_code=400, detail="text or stored_filename required")

    # --- Clause segmentation ---
    try:
        clauses = FileProcessor.preprocess_text(text)
    except Exception:
        # fallback: split on double newlines or sentences if blank
        clauses = [c.strip() for c in re.split(r'\n{2,}', text) if c.strip()]

    if not clauses:
        clauses = [text]

    results: List[Dict[str, Any]] = []
    total_risk = 0.0
    MAX_CLAUSES = 50

    # --- Analyze each clause ---
    for clause in clauses[:MAX_CLAUSES]:
        # clause type + confidence
        try:
            clause_type, confidence = legal_bert.predict_clause_type(clause)
        except Exception:
            clause_type, confidence = "general", 0.0

        # risk analysis
        try:
            risk_data = risk_engine.analyze_risk_with_statutes(clause, clause_type)
        except Exception:
            risk_data = {
                "risk_level": "low",
                "risk_score": 0.0,
                "violations": [],
                "compliance_issues": [],
                "legal_references": [],
                "pattern_violations": []
            }

        # legal references (engine helper)
        try:
            legal_refs = risk_engine.get_legal_references(clause, clause_type)
        except Exception:
            legal_refs = risk_data.get("legal_references", [])

        # summary
        try:
            summary = legal_bert.generate_dynamic_summary(clause)
        except Exception:
            summary = (clause[:300] + "...") if len(clause) > 300 else clause

        # recommendations
        try:
            recs = risk_engine.generate_dynamic_recommendations(clause, risk_data, legal_refs)
        except Exception:
            recs = ["Review clause carefully; automated recommendations unavailable."]

        # decide if clause is risky (>= 20%)
        risk_score = float(risk_data.get("risk_score", 0.0))
        is_risky = risk_score >= 0.20

        # generate improved clause only for risky ones
        improved_clause = None
        if is_risky:
            try:
                improved_clause = generate_improved_clause(clause, clause_type)
            except Exception:
                improved_clause = None

        results.append({
            "clause_text": clause,                       # original full clause
            "clause_type": clause_type,
            "confidence": float(confidence),
            "risk_level": risk_data.get("risk_level", "low"),
            "risk_score": risk_score,
            "summary": summary,
            "violations": risk_data.get("violations", []),
            "compliance_issues": risk_data.get("compliance_issues", []),
            "recommendations": recs,
            "legal_references": legal_refs,
            "improved_clause": improved_clause
        })

        total_risk += risk_score

    # --- Aggregation logic (improved weighting) ---
    scores = [r.get("risk_score", 0.0) for r in results] or [0.0]
    max_score = max(scores)
    avg_all = sum(scores) / len(scores) if scores else 0.0
    TOP_K = 5
    top_k_scores = sorted(scores, reverse=True)[:TOP_K]
    avg_top = sum(top_k_scores) / len(top_k_scores) if top_k_scores else 0.0

    # Combine into overall risk (weighted)
    overall_risk_score = (0.50 * max_score) + (0.35 * avg_top) + (0.15 * avg_all)
    overall_risk_score = max(0.0, min(1.0, overall_risk_score))

    # risk level thresholds (tuneable)
    if overall_risk_score >= 0.65:
        overall_risk_level = "High"
    elif overall_risk_score >= 0.35:
        overall_risk_level = "Medium"
    else:
        overall_risk_level = "Low"

    # --- Prepare risky clause filtering & summaries ---
    RISK_THRESHOLD = 0.2  # 20%
    top_sorted = sorted(results, key=lambda x: x["risk_score"], reverse=True)
    risky_clauses = [r for r in top_sorted if r["risk_score"] >= RISK_THRESHOLD]

    # detailed_summary: show only when no risky clauses (per requirement)
    if not risky_clauses:
        detailed_summary = (
            "✅ No significant risks detected.\n"
            "All clauses in this document have risk scores below 20%.\n"
            "The document appears legally safe and compliant.\n"
        )
    else:
        detailed_summary = ""  # hide when risky clauses exist

    # --- Build recommendations (deduped) only from risky clauses ---
    recs_acc: List[str] = []
    law_set = set()
    for r in risky_clauses:
        for rec in r.get("recommendations", []):
            if isinstance(rec, str):
                # skip lines that are explicit law markers (we'll collect laws separately)
                if rec.strip().startswith("📘") or rec.strip().lower().startswith("reference") or rec.strip().lower().startswith("- reference"):
                    continue
                if rec.strip() and rec.strip() not in recs_acc:
                    recs_acc.append(rec.strip())
        for ref in r.get("legal_references", []):
            if isinstance(ref, dict):
                statute = ref.get("statute", "").strip()
                section = ref.get("section", "").strip()
                if statute:
                    if section and section not in statute:
                        law_set.add(f"{statute} – {section}")
                    else:
                        law_set.add(statute)
            elif isinstance(ref, str):
                law_set.add(ref.strip())

    if not recs_acc and not risky_clauses:
        recs_acc = [
            "✅ No significant risks detected.",
            "The document appears legally compliant and safe."
        ]

    final_recommendations = []
    for r in recs_acc:
        if r not in final_recommendations:
            final_recommendations.append(r)
    final_recommendations = final_recommendations[:20]

    relevant_laws = sorted(list(law_set))

    # --- Produce risky_summary (plain text) for quick frontend display ---
    risky_summary = ""
    if risky_clauses:
        risky_summary += "⚠️ Top Risky Clauses\n(Only clauses with risk ≥ 20% are shown)\n\n"
        for idx, r in enumerate(risky_clauses[:10], start=1):
            risk_pct = round(r.get("risk_score", 0.0) * 100)
            risky_summary += (
                f"{idx}. {r.get('clause_type','Unknown').upper()} — "
                f"{r.get('risk_level','Unknown').upper()} ({risk_pct}%)\n"
                f"{r.get('summary')}\n\n"
            )

    # --- Prepare risky_clauses_out: include both original + improved clause ---
    risky_clauses_out = []
    for r in risky_clauses[:10]:
        risky_clauses_out.append({
            "clause_type": r.get("clause_type"),
            "risk_level": r.get("risk_level"),
            "risk_score": float(r.get("risk_score", 0.0)),
            "summary": r.get("summary"),
            "clause_text": r.get("clause_text"),
            "improved_clause": r.get("improved_clause"),
            "recommendations": r.get("recommendations", [])[:6],
            "legal_references": r.get("legal_references", [])[:6]
        })

    detected_risks = [r.get("clause_type", "unknown") for r in results]

    # --- Generate user-facing message ---
    if risky_clauses:
        if overall_risk_level == "High":
            message = "🚨 Several problematic clauses detected. Strongly consider revisions."
        elif overall_risk_level == "Medium":
            message = "⚠️ Some clauses may pose moderate risk. Review recommended."
        else:
            message = "⚠️ Risky clauses detected. Please review recommendations."
    else:
        message = "✅ No significant risks found. The document appears safe."

    # --- Document-level summary (try best-effort) ---
    try:
        document_summary = legal_bert.generate_dynamic_summary(text, max_length=150)
    except Exception:
        document_summary = " ".join([r["summary"] for r in results[:5]])

    # Optional debug aggregation info (handy for frontend tuning)
    aggregation_debug = {
        "max_clause_score": round(max_score * 100),
        "avg_top_k_score": round(avg_top * 100),
        "avg_all_score": round(avg_all * 100),
        "combined_overall_score": round(overall_risk_score * 100),
    }
    # Build corrected document for display
    corrected_document = build_corrected_document(clauses, results, threshold=0.20)

    # --- Final response ---
    return {
        "summary": document_summary,
        "overall_risk_score": round(overall_risk_score, 4),   # 0..1 float
        "overall_risk_level": overall_risk_level,
        "detailed_summary": detailed_summary,
        "risky_clauses": risky_clauses_out,
        "risky_summary": risky_summary,
        "detected_risks": detected_risks,
        "recommendations": final_recommendations,
        "relevant_laws": relevant_laws,
        "corrected_document": corrected_document,
        "count": len(results),
        "message": message,
        "has_corrections": any(
    r.get("improved_clause") and r.get("improved_clause") != r.get("clause_text")
    for r in risky_clauses
)
        ,
        "original_clauses": clauses,
        "aggregation_debug": aggregation_debug
    }

from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

@app.post("/download-docx")
async def download_docx(payload: dict, user=Depends(get_current_user)):
    corrected_text = payload.get("corrected_text", "")
    filename = payload.get("filename", "Corrected_Document.docx")

    if not corrected_text:
        raise HTTPException(status_code=400, detail="Missing corrected_text")

    output_path = os.path.join(UPLOAD_DIR, filename)

    # ---- Create DOCX with preserved formatting ----
    doc = Document()

    # Standard legal doc formatting
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    paragraphs = corrected_text.split("\n")

    for line in paragraphs:
        if line.strip() == "":
            # Blank line = paragraph break
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph()
        r = p.add_run(line)

        # Legal document paragraph styling
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    doc.save(output_path)

    # ---- Return file ----
    response = FileResponse(
        output_path,
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        filename=filename,
    )

    # Required CORS headers
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Expose-Headers"] = "Content-Disposition"

    return response
